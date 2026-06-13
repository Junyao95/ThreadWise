# document_processor.py
import io
import pypdf
import msoffcrypto
import pandas as pd

def decrypt_pdf(content_bytes: bytes, password: str) -> str:
    """
    Decrypt a password-protected PDF and extract all text.
    """
    reader = pypdf.PdfReader(io.BytesIO(content_bytes))
    if reader.is_encrypted:
        # decrypt with provided password
        reader.decrypt(password)
    # Extract text from all pages
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

# document_processor.py (add to existing)

def decrypt_word(content_bytes: bytes, password: str) -> str:
    """
    Decrypt a password-protected Word (.docx) document and extract text.
    """
    import docx
    # Use msoffcrypto to decrypt
    decrypted_stream = io.BytesIO()
    with io.BytesIO(content_bytes) as encrypted_stream:
        office_file = msoffcrypto.OfficeFile(encrypted_stream)
        office_file.load_key(password=password)
        office_file.decrypt(decrypted_stream)

    # Now read the decrypted docx
    decrypted_stream.seek(0)
    doc = docx.Document(decrypted_stream)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

# document_processor.py (add to existing)

def decrypt_excel(content_bytes: bytes, password: str) -> str:
    """
    Decrypt a password-protected Excel (.xlsx) document and extract data.
    """
    decrypted_stream = io.BytesIO()
    with io.BytesIO(content_bytes) as encrypted_stream:
        office_file = msoffcrypto.OfficeFile(encrypted_stream)
        office_file.load_key(password=password)
        office_file.decrypt(decrypted_stream)

    decrypted_stream.seek(0)
    # Read all sheets into a single text summary
    all_sheets_text = []
    xls = pd.ExcelFile(decrypted_stream)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(decrypted_stream, sheet_name=sheet_name)
        all_sheets_text.append(f"Sheet: {sheet_name}\n{df.to_string()}\n")
        # Reset stream position for next sheet
        decrypted_stream.seek(0)

    return "\n".join(all_sheets_text)

def process_attachment(content_bytes: bytes, filename: str, password: str = None) -> str:
    """
    Route attachment to appropriate decryption and extraction function.
    Returns extracted text content as a string.
    """
    if filename.lower().endswith('.pdf'):
        if password:
            return decrypt_pdf(content_bytes, password)
        else:
            # Attempt to read as unencrypted PDF
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

    elif filename.lower().endswith('.docx') or filename.lower().endswith('.doc'):
        if password:
            return decrypt_word(content_bytes, password)
        else:
            # Attempt unencrypted read
            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            return "\n".join([p.text for p in doc.paragraphs])

    elif filename.lower().endswith('.xlsx') or filename.lower().endswith('.xls'):
        if password:
            return decrypt_excel(content_bytes, password)
        else:
            # Attempt unencrypted read
            return pd.read_excel(io.BytesIO(content_bytes)).to_string()

    else:
        # Unsupported file type – return placeholder
        return f"[Unsupported file type: {filename}]"